"""
Knowledge Service

This service handles all knowledge-related business logic including knowledge
configuration, file uploads, vector database operations, and component discovery.
"""

import os
import sys
from datetime import datetime
from typing import List, Optional, Dict, Any
from fastapi import HTTPException, status, UploadFile

from ..db import get_collections
from ..utils.log import logger
from src.utils.component_discovery import discover_components, get_detailed_class_info


class KnowledgeService:
    """Service for managing knowledge configurations and operations"""
    
    @staticmethod
    def _get_knowledge_config_collection():
        """Get the knowledge config collection"""
        return get_collections()["knowledgeConfig"]
    
    @staticmethod
    async def validate_tenant_access(user: dict) -> str:
        """Validate tenant access and return tenant_id"""
        tenant_id = user.get("tenantId")
        if not tenant_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User tenant information missing. Please re-login."
            )
        return tenant_id
    
    @classmethod
    async def discover_chunking_components(cls) -> List[Dict[str, Any]]:
        """Discover available chunking components"""
        logger.info("[KNOWLEDGE] Discovering chunking components")
        
        try:
            components = discover_components("ai.document.chunking")
            result = []
            
            for comp in components:
                try:
                    module_path = f"ai.document.chunking.{comp}"
                    class_info = get_detailed_class_info(module_path, comp)
                    
                    result.append({
                        "name": comp,
                        "module": module_path,
                        "description": class_info.get("description", ""),
                        "parameters": class_info.get("parameters", {}),
                        "examples": class_info.get("examples", {})
                    })
                except Exception as e:
                    logger.warning(f"[KNOWLEDGE] Failed to introspect {comp}: {e}")
                    result.append({
                        "name": comp,
                        "module": f"ai.document.chunking.{comp}",
                        "description": "Component introspection failed",
                        "parameters": {},
                        "examples": {}
                    })
            
            logger.info(f"[KNOWLEDGE] Found {len(result)} chunking components")
            return result
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to discover components: {e}")
            raise HTTPException(status_code=500, detail="Failed to discover chunking components")
    
    @classmethod
    async def get_component_details(cls, component_name: str) -> Dict[str, Any]:
        """Get detailed information about a specific component"""
        logger.info(f"[KNOWLEDGE] Getting details for component: {component_name}")
        
        try:
            module_path = f"ai.document.chunking.{component_name}"
            class_info = get_detailed_class_info(module_path, component_name)
            
            return {
                "name": component_name,
                "module": module_path,
                "description": class_info.get("description", ""),
                "parameters": class_info.get("parameters", {}),
                "examples": class_info.get("examples", {})
            }
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to get component details for {component_name}: {e}")
            raise HTTPException(status_code=404, detail=f"Component '{component_name}' not found")
    
    @classmethod
    async def list_knowledge_configs(cls, user: dict) -> List[Dict[str, Any]]:
        """List knowledge configurations for current tenant"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId")
        
        logger.info(f"[KNOWLEDGE] Listing configs for tenant: {tenant_id}, user: {user_id}")
        
        try:
            cursor = cls._get_knowledge_config_collection().find({
                "tenantId": tenant_id,
                "userId": user_id
            }).sort("name", 1)
            
            docs = await cursor.to_list(length=None)
            
            configs = []
            for doc in docs:
                config = {
                    "id": str(doc["_id"]),
                    "name": doc.get("name"),
                    "description": doc.get("description", ""),
                    "chunker": doc.get("chunker", {}),
                    "created_at": doc.get("created_at"),
                    "updated_at": doc.get("updated_at"),
                    "file_count": doc.get("file_count", 0)
                }
                configs.append(config)
            
            logger.info(f"[KNOWLEDGE] Found {len(configs)} configurations")
            return configs
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to list configs: {e}")
            raise HTTPException(status_code=500, detail="Failed to retrieve knowledge configurations")
    
    @classmethod
    async def get_knowledge_config(cls, collection_name: str, user: dict) -> Dict[str, Any]:
        """Get specific knowledge configuration"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId")
        
        try:
            doc = await cls._get_knowledge_config_collection().find_one({
                "tenantId": tenant_id,
                "userId": user_id,
                "name": collection_name
            })
            
            if not doc:
                raise HTTPException(status_code=404, detail="Knowledge configuration not found")
            
            config = {
                "id": str(doc["_id"]),
                "name": doc.get("name"),
                "description": doc.get("description", ""),
                "chunker": doc.get("chunker", {}),
                "created_at": doc.get("created_at"),
                "updated_at": doc.get("updated_at"),
                "file_count": doc.get("file_count", 0)
            }
            
            return config
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to get config {collection_name}: {e}")
            raise HTTPException(status_code=500, detail="Failed to retrieve knowledge configuration")
    
    @classmethod
    async def create_knowledge_config(cls, config_data: Dict[str, Any], user: dict) -> Dict[str, str]:
        """Create new knowledge configuration"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId")
        
        name = config_data.get("name", "").strip()
        if not name:
            raise HTTPException(status_code=400, detail="Configuration name is required")
        
        logger.info(f"[KNOWLEDGE] Creating config '{name}' for tenant: {tenant_id}, user: {user_id}")
        
        try:
            # Check if config already exists
            existing = await cls._get_knowledge_config_collection().find_one({
                "tenantId": tenant_id,
                "userId": user_id,
                "name": name
            })
            
            if existing:
                raise HTTPException(status_code=409, detail="Configuration with this name already exists")
            
            # Create configuration record
            record = {
                "tenantId": tenant_id,
                "userId": user_id,
                "name": name,
                "description": config_data.get("description", ""),
                "chunker": config_data.get("chunker", {}),
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                "file_count": 0
            }
            
            await cls._get_knowledge_config_collection().insert_one(record)
            
            # Initialize vector collection
            await cls._initialize_vector_collection(tenant_id, user_id, name)
            
            logger.info(f"[KNOWLEDGE] Successfully created config '{name}'")
            return {"message": "Knowledge configuration created", "name": name}
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to create config '{name}': {e}")
            raise HTTPException(status_code=500, detail="Failed to create knowledge configuration")
    
    @classmethod
    async def update_knowledge_config(cls, collection_name: str, config_data: Dict[str, Any], user: dict) -> Dict[str, str]:
        """Update existing knowledge configuration"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId")
        
        logger.info(f"[KNOWLEDGE] Updating config '{collection_name}' for tenant: {tenant_id}")
        
        try:
            update_data = {
                "updated_at": datetime.utcnow()
            }
            
            # Only update allowed fields
            allowed_fields = ["description", "chunker"]
            for field in allowed_fields:
                if field in config_data:
                    update_data[field] = config_data[field]
            
            result = await cls._get_knowledge_config_collection().update_one(
                {
                    "tenantId": tenant_id,
                    "userId": user_id,
                    "name": collection_name
                },
                {"$set": update_data}
            )
            
            if result.matched_count == 0:
                raise HTTPException(status_code=404, detail="Knowledge configuration not found")
            
            logger.info(f"[KNOWLEDGE] Successfully updated config '{collection_name}'")
            return {"message": "Knowledge configuration updated"}
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to update config '{collection_name}': {e}")
            raise HTTPException(status_code=500, detail="Failed to update knowledge configuration")
    
    @classmethod
    async def delete_knowledge_config(cls, collection_name: str, user: dict) -> Dict[str, str]:
        """Delete knowledge configuration and associated vector collection"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId")
        
        logger.info(f"[KNOWLEDGE] Deleting config '{collection_name}' for tenant: {tenant_id}")
        
        try:
            # Delete configuration
            result = await cls._get_knowledge_config_collection().delete_one({
                "tenantId": tenant_id,
                "userId": user_id,
                "name": collection_name
            })
            
            if result.deleted_count == 0:
                raise HTTPException(status_code=404, detail="Knowledge configuration not found")
            
            # Delete vector collection
            await cls._delete_vector_collection(tenant_id, user_id, collection_name)
            
            logger.info(f"[KNOWLEDGE] Successfully deleted config '{collection_name}'")
            return {"message": f"Knowledge configuration '{collection_name}' deleted"}
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to delete config '{collection_name}': {e}")
            raise HTTPException(status_code=500, detail="Failed to delete knowledge configuration")
    
    @classmethod
    async def _initialize_vector_collection(cls, tenant_id: str, user_id: str, collection: str):
        """Initialize a new vector database collection for this knowledge collection"""
        try:
            from ai.vectordb.qdrant import Qdrant
            from ai.embedder.openai import OpenAIEmbedder
            
            # Create collection name with user isolation: collectionname_userid
            vector_collection_name = f"{collection}_{user_id}"
            
            # Initialize vector database
            vector_db = Qdrant(
                collection=vector_collection_name,
                host=os.getenv("QDRANT_HOST", "localhost"),
                port=int(os.getenv("QDRANT_PORT", "8805")),
                https=os.getenv("QDRANT_HTTPS", "false").lower() == "true",
                api_key=os.getenv("QDRANT_API_KEY", None),
                embedder=OpenAIEmbedder()
            )
            
            # Create the collection in Qdrant
            vector_db.create()
            logger.info(f"[KNOWLEDGE] Initialized vector collection: {vector_collection_name}")
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to initialize vector collection {collection}: {e}")
            raise
    
    @classmethod
    async def _delete_vector_collection(cls, tenant_id: str, user_id: str, collection: str):
        """Delete a vector database collection"""
        try:
            from ai.vectordb.qdrant import Qdrant
            
            # Create collection name with user isolation: collectionname_userid
            vector_collection_name = f"{collection}_{user_id}"
            
            # Initialize vector database client
            vector_db = Qdrant(
                collection=vector_collection_name,
                host=os.getenv("QDRANT_HOST", "localhost"),
                port=int(os.getenv("QDRANT_PORT", "8805")),
                https=os.getenv("QDRANT_HTTPS", "false").lower() == "true",
                api_key=os.getenv("QDRANT_API_KEY", None),
            )
            
            # Delete the collection
            vector_db.delete()
            logger.info(f"[KNOWLEDGE] Deleted vector collection: {vector_collection_name}")
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to delete vector collection {collection}: {e}")
            # Don't raise exception for cleanup operations
    
    @classmethod  
    async def _index_file_to_vector_db(cls, tenant_id: str, user_id: str, collection: str, filename: str, content: bytes, object_key: str):
        """Index a file's content to the vector database"""
        try:
            from ai.vectordb.qdrant import Qdrant
            from ai.embedder.openai import OpenAIEmbedder
            from ai.document import Document
            from ai.document.chunking.fixed import FixedChunker
            
            # Create collection name with user isolation: collectionname_userid
            vector_collection_name = f"{collection}_{user_id}"
            
            # Initialize vector database
            vector_db = Qdrant(
                collection=vector_collection_name,
                host=os.getenv("QDRANT_HOST", "localhost"),
                port=int(os.getenv("QDRANT_PORT", "8805")),
                https=os.getenv("QDRANT_HTTPS", "false").lower() == "true",
                api_key=os.getenv("QDRANT_API_KEY", None),
                embedder=OpenAIEmbedder()
            )
            
            # Extract text from file
            text_content = await cls._extract_text_from_file(filename, content)
            
            # Create a document
            document = Document(
                name=filename,
                id=object_key,
                content=text_content,
                meta={"source": filename, "object_key": object_key}
            )
            
            # Chunk the document
            chunker = FixedChunker(chunk_size=1024, overlap=64)
            chunks = chunker.chunk([document])
            
            # Index the chunks
            vector_db.insert(chunks)
            logger.info(f"[KNOWLEDGE] Indexed {len(chunks)} chunks from {filename} to collection {vector_collection_name}")
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to index file {filename} to vector collection {collection}: {e}")
            raise
    
    @classmethod
    async def _extract_text_from_file(cls, filename: str, content: bytes) -> str:
        """Extract text content from uploaded file"""
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'txt':
                return content.decode('utf-8')
            elif file_extension == 'pdf':
                import PyPDF2
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
            elif file_extension in ['doc', 'docx']:
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            elif file_extension == 'md':
                return content.decode('utf-8')
            else:
                # Try to decode as text
                return content.decode('utf-8')
                
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to extract text from {filename}: {e}")
            # Return empty string if extraction fails
            return ""
    
    @classmethod
    def _get_minio_client(cls):
        """Get MinIO client for file storage"""
        from minio import Minio
        
        return Minio(
            f"{os.getenv('MINIO_HOST', 'localhost')}:{os.getenv('MINIO_PORT', '9000')}",
            access_key=os.getenv('MINIO_ACCESS_KEY', 'minioadmin'),
            secret_key=os.getenv('MINIO_SECRET_KEY', 'minioadmin'),
            secure=os.getenv('MINIO_SECURE', 'false').lower() == 'true'
        )
    
    @classmethod
    async def upload_files(cls, user: dict, files: List[UploadFile], knowledge_prefix: str):
        """Upload files to MinIO and index them in vector database"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id")
        
        if not user_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User ID missing from token"
            )
        
        # Validate knowledge prefix exists
        config = await cls.get_knowledge_config(user, knowledge_prefix)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Knowledge prefix '{knowledge_prefix}' not found"
            )
        
        logger.info(f"[KNOWLEDGE] Uploading {len(files)} files for prefix: {knowledge_prefix}")
        
        try:
            minio_client = cls._get_minio_client()
            bucket_name = "uploads"
            
            # Ensure bucket exists
            if not minio_client.bucket_exists(bucket_name):
                minio_client.make_bucket(bucket_name)
            
            uploaded_files = []
            
            for file in files:
                # Read file content
                content = await file.read()
                
                # Create object key: uploads/{tenant_id}/{user_id}/{prefix}/{filename}
                object_key = f"uploads/{tenant_id}/{user_id}/{knowledge_prefix}/{file.filename}"
                
                # Upload to MinIO
                from io import BytesIO
                minio_client.put_object(
                    bucket_name,
                    object_key,
                    BytesIO(content),
                    length=len(content),
                    content_type=file.content_type or 'application/octet-stream'
                )
                
                # Index to vector database if configured
                try:
                    await cls._index_file_to_vector_db(
                        tenant_id, user_id, knowledge_prefix,
                        file.filename, content, object_key
                    )
                except Exception as e:
                    logger.error(f"[KNOWLEDGE] Failed to index {file.filename}: {e}")
                    # Continue with upload even if indexing fails
                
                uploaded_files.append({
                    "filename": file.filename,
                    "object_key": object_key,
                    "size": len(content),
                    "content_type": file.content_type
                })
                
                logger.info(f"[KNOWLEDGE] Uploaded and indexed: {file.filename}")
            
            return {
                "message": f"Successfully uploaded {len(uploaded_files)} files",
                "files": uploaded_files
            }
            
        except Exception as e:
            logger.error(f"[KNOWLEDGE] File upload failed: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"File upload failed: {str(e)}"
            )
    
    @classmethod
    async def list_categories(cls, user: dict) -> List[str]:
        """List knowledge categories for current tenant"""
        tenant_id = await cls.validate_tenant_access(user)
        collection = cls._get_knowledge_config_collection()
        
        try:
            categories = await collection.distinct("category", {"tenantId": tenant_id})
            return sorted(categories)
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to list categories: {e}")
            raise HTTPException(status_code=500, detail="Failed to list categories")
    
    @classmethod
    async def introspect_component(cls, module_name: str) -> Dict[str, Any]:
        """Introspect a chunking component to get its parameters"""
        logger.info(f"[KNOWLEDGE] Introspecting component: {module_name}")
        
        try:
            class_info = get_detailed_class_info(module_name, module_name.split('.')[-1])
            return {
                "module": module_name,
                "description": class_info.get("description", ""),
                "parameters": class_info.get("parameters", {}),
                "examples": class_info.get("examples", {})
            }
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to introspect {module_name}: {e}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Failed to introspect module '{module_name}'"
            )
    
    @classmethod
    async def upsert_knowledge_config(cls, user: dict, payload: dict) -> dict:
        """Create or update a knowledge prefix configuration"""
        from datetime import datetime
        
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId") or "unknown"
        
        name = (payload.get("name") or "").strip()
        if not name:
            raise HTTPException(status_code=400, detail="name is required")

        record = {
            "name": name,
            "category": payload.get("category", ""),
            "chunk_strategy": payload.get("chunk_strategy", ""),
            "chunk_strategy_params": payload.get("chunk_strategy_params", {}),
            "chunk_size": payload.get("chunk_size", 5000),
            "chunk_overlap": payload.get("chunk_overlap", 0),
            "add_context": payload.get("add_context", True),
            "search_knowledge": payload.get("search_knowledge", True),
            "files": payload.get("files", []),
            "tenantId": tenant_id,
            "userId": user_id,
            "updated_at": datetime.utcnow(),
        }

        collection = cls._get_knowledge_config_collection()
        
        # Create vs update
        existing = await collection.find_one({"name": name, "tenantId": tenant_id})
        if existing:
            await collection.update_one({"_id": existing["_id"]}, {"$set": record})
            action = "updated"
        else:
            record["created_at"] = datetime.utcnow()
            await collection.insert_one(record)
            action = "created"
            
            # Initialize vector collection for new knowledge config
            try:
                await cls._initialize_vector_collection(tenant_id, user_id, name)
            except Exception as e:
                logger.error(f"[KNOWLEDGE] Failed to initialize vector collection for {name}: {e}")
                # Don't fail the entire operation for vector DB issues

        return {"message": f"Prefix {action}", "name": name}
    
    @classmethod
    async def delete_knowledge_config(cls, user: dict, name: str) -> dict:
        """Delete a knowledge prefix configuration"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id") or user.get("userId") or "unknown"
        
        collection = cls._get_knowledge_config_collection()
        
        # Check if config exists
        existing = await collection.find_one({"name": name, "tenantId": tenant_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Prefix not found")
        
        # Delete from MongoDB
        result = await collection.delete_one({"name": name, "tenantId": tenant_id})
        
        # Delete vector collection
        try:
            await cls._delete_vector_collection(tenant_id, user_id, name)
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Failed to delete vector collection for {name}: {e}")
            # Don't fail the entire operation for vector DB cleanup issues
        
        return {"message": f"Prefix '{name}' deleted"}
    
    @classmethod
    async def get_defaults(cls, user: dict) -> Dict[str, Any]:
        """Return global defaults for knowledge chunking configuration"""
        defaults = {
            "chunk_size": int(os.getenv("KNOWLEDGE_CHUNK_SIZE", "5000")),
            "chunk_overlap": int(os.getenv("KNOWLEDGE_CHUNK_OVERLAP", "0")),
            "add_context": os.getenv("KNOWLEDGE_ADD_CONTEXT", "true").lower() == "true",
            "search_knowledge": os.getenv("KNOWLEDGE_SEARCH", "true").lower() == "true",
        }
        return {"defaults": defaults}
    
    @classmethod
    async def list_categories(cls, user: dict) -> Dict[str, List[str]]:
        """Return distinct knowledge categories for the user's tenant"""
        tenant_id = await cls.validate_tenant_access(user)
        
        collection = cls._get_knowledge_config_collection()
        cats = await collection.distinct("category", {"tenantId": tenant_id})
        # filter out empty
        categories = [c for c in cats if c]
        return {"categories": categories}
    
    @classmethod
    async def list_collections(cls, user: dict) -> Dict[str, List[str]]:
        """List knowledge collections for current tenant"""
        tenant_id = await cls.validate_tenant_access(user)
        
        collection = cls._get_knowledge_config_collection()
        cursor = collection.find({"tenantId": tenant_id}, {"collection": 1}).sort("collection", 1)
        docs = await cursor.to_list(length=None)
        collections_list = sorted({d.get("collection") for d in docs if d.get("collection")})
        return {"collections": list(collections_list)}
    
    @classmethod
    async def get_collection(cls, collection_name: str, user: dict) -> Dict[str, Any]:
        """Get collection details with files from MinIO"""
        tenant_id = await cls.validate_tenant_access(user)
        
        collection = cls._get_knowledge_config_collection()
        doc = await collection.find_one({"tenantId": tenant_id, "collection": collection_name})
        if not doc:
            raise HTTPException(status_code=404, detail="collection not found")

        # Load files from MinIO
        try:
            from .file_service import FileService
            
            # If ownerId saved in doc use it, else current user
            owner_id = doc.get("ownerId") or user.get("id")
            path_prefix = f"uploads/{tenant_id}/{owner_id}/{collection_name}/"
            
            file_names = await FileService.list_files_by_prefix(path_prefix)
            doc["files"] = file_names
        except Exception as e:
            logger.warning(f"[KNOWLEDGE] MinIO listing failed for collection {collection_name}: {e}")
            doc["files"] = []

        # shape compatible with UI
        owner_id = doc.get("ownerId") or user.get("id")
        return {
            "collection": doc.get("collection"),
            "vector_collection": doc.get("vector_collection", f"{doc.get('collection', '')}_user@{owner_id}"),
            "category": doc.get("category", ""),
            "chunk": doc.get("chunk", {}),
            "created_at": doc.get("created_at"),
            "updated_at": doc.get("updated_at"),
            "files": [f"{doc.get('collection')}/{fn}" for fn in doc.get("files", [])],
            "doc_counts": doc.get("doc_counts", {}),
            "files_by_type": doc.get("files_by_type", {}),
            "file_count": len(doc.get("files", [])),
        }
    
    @classmethod
    async def save_collection(cls, payload: dict, user: dict) -> Dict[str, Any]:
        """Create or update a knowledge collection configuration in MongoDB and trigger indexing"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id")
        if not user_id or not user_id.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User ID missing. Please re-login."
            )

        collection_name = payload.get("collection") or payload.get("name")
        if not collection_name or not collection_name.strip():
            raise HTTPException(status_code=400, detail="collection is required")

        # Clean collection name (remove special characters that might cause issues)
        collection_name = collection_name.strip()

        overwrite = bool(payload.get("overwrite"))
        now_ms = int(datetime.utcnow().timestamp() * 1000)

        doc = await cls._get_knowledge_config_collection().find_one({"tenantId": tenant_id, "collection": collection_name})
        if doc and not overwrite:
            return {"exists": True, "message": "collection exists"}

        # Create vector collection name for storage in MongoDB
        vector_collection_name = f"{collection_name}_user@{user_id}"

        record = {
            "tenantId": tenant_id,
            "ownerId": user_id,
            "collection": collection_name,
            "vector_collection": vector_collection_name,  # Store vector collection name
            "category": payload.get("category", ""),
            "chunk": payload.get("chunk", {
                "chunk_size": payload.get("chunk_size"),
                "chunk_overlap": payload.get("chunk_overlap"),
                "add_context": payload.get("add_context"),
                "search_knowledge": payload.get("search_knowledge"),
            }),
            "updated_at": now_ms,
        }
        if not doc:
            record["created_at"] = now_ms

        # Upsert
        await cls._get_knowledge_config_collection().update_one(
            {"tenantId": tenant_id, "collection": collection_name},
            {"$set": record},
            upsert=True
        )

        # If this is a new collection creation, initialize vector DB collection
        if not doc:
            try:
                await cls._initialize_vector_collection(tenant_id, user_id, collection_name)
            except Exception as e:
                logger.error(f"[KNOWLEDGE] Failed to initialize vector collection: {e}")
                # Don't fail the entire operation for vector DB issues

        return {"ok": True, "collection": collection_name, "vector_collection": vector_collection_name}

    @classmethod  
    async def delete_collection(cls, collection_name: str, user: dict) -> Dict[str, Any]:
        """Delete knowledge collection configuration and vector collection"""
        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id")
        if not user_id or not user_id.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User ID missing. Please re-login."
            )
        
        res = await cls._get_knowledge_config_collection().delete_one({"tenantId": tenant_id, "collection": collection_name})
        if res.deleted_count == 0:
            raise HTTPException(status_code=404, detail="collection not found")
        
        # Also try to delete the vector DB collection
        try:
            await cls._delete_vector_collection(tenant_id, user_id, collection_name)
        except Exception as e:
            logger.warning(f"[KNOWLEDGE] Failed to delete vector collection: {e}")
        
        return {"deleted": True, "collection": collection_name, "vector_collection": f"{collection_name}_user@{user_id}"}
    
    @classmethod
    async def upload_knowledge_files(cls, collection_name: str, files: List[UploadFile], user: dict) -> Dict[str, Any]:
        """Upload files to MinIO under uploads/tenantId/userId/collection/ and index them"""
        from io import BytesIO
        
        if not files:
            raise HTTPException(status_code=400, detail="No files provided")

        tenant_id = await cls.validate_tenant_access(user)
        user_id = user.get("id")
        if not user_id or not user_id.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User ID missing. Cannot upload files without user ID."
            )

        if not collection_name or not collection_name.strip():
            raise HTTPException(status_code=400, detail="Collection name is required")

        # Clean the collection name
        collection_name = collection_name.strip()

        # Constants
        ALLOWED_TYPES = {
            'application/pdf',
            'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain', 'text/csv', 'application/json',
            'text/x-python', 'application/javascript', 'text/javascript',
        }
        MAX_FILE_SIZE = 200 * 1024 * 1024

        try:
            from .file_service import FileService
            
            results = []
            base_prefix = f"uploads/{tenant_id}/{user_id}/{collection_name}"

            for f in files:
                if f.content_type and f.content_type not in ALLOWED_TYPES:
                    raise HTTPException(status_code=400, detail=f"File type not allowed: {f.filename}")

                content = await f.read()
                if len(content) > MAX_FILE_SIZE:
                    raise HTTPException(status_code=400, detail=f"File too large: {f.filename}")

                ts = int(datetime.utcnow().timestamp() * 1000)
                object_name = f"{base_prefix}/{ts}-{f.filename}"
                
                # Upload file
                await FileService.upload_file_content(object_name, content, f.content_type)

                results.append({
                    "filename": f.filename,
                    "size": len(content),
                    "key": object_name,
                    "content": content,  # Keep content for indexing
                })

            # Index uploaded files to vector database
            indexed_files = []
            for file_result in results:
                try:
                    await cls._index_file_to_vector_db(
                        tenant_id=tenant_id,
                        user_id=user_id,
                        collection=collection_name,
                        filename=file_result["filename"],
                        content=file_result["content"],
                        object_key=file_result["key"]
                    )
                    indexed_files.append(file_result["filename"])
                except Exception as e:
                    logger.error(f"[KNOWLEDGE] Failed to index file {file_result['filename']}: {e}")

            # Update the MongoDB collection with file information
            vector_collection_name = f"{collection_name}_user@{user_id}"
            
            await cls._get_knowledge_config_collection().update_one(
                {"tenantId": tenant_id, "collection": collection_name},
                {
                    "$addToSet": {"files": {"$each": [r["filename"] for r in results]}},
                    "$set": {
                        "updated_at": int(datetime.utcnow().timestamp() * 1000),
                        "vector_collection": vector_collection_name,
                    }
                },
                upsert=True
            )

            # Remove content from results before returning (too large for response)
            for result in results:
                result.pop("content", None)

            return {
                "uploaded": results, 
                "collection": collection_name,
                "vector_collection": vector_collection_name,
                "indexed_files": indexed_files,
                "tenant_id": tenant_id,
                "user_id": user_id
            }

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"[KNOWLEDGE] Upload failed: {e}")
            raise HTTPException(status_code=500, detail="Upload failed")
    
    @classmethod
    async def get_diagnostics(cls) -> Dict[str, Any]:
        """Get knowledge service diagnostics"""
        try:
            from minio import Minio
            minio_status = "ok"
        except Exception as e:
            minio_status = f"missing: {e}"
        
        return {
            "minioImport": minio_status,
            "endpoint": os.getenv("MINIO_ENDPOINT", "127.0.0.1:8803"),
            "bucket": os.getenv("MINIO_BUCKET", "hcp"),
        }
